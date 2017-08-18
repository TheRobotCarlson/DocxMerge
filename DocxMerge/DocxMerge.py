from difflib import ndiff, SequenceMatcher
import string
from docx import Document as RetDoc
from shutil import copyfile
from docx.document import Document


def get_diffs_from_seq(text_a, text_b):
    sm = SequenceMatcher(isjunk=lambda x: x in " \t\n", a=text_a, b=text_b)
    opcodes = sm.get_opcodes()

    replacements = {}

    for tag, i1, i2, j1, j2 in opcodes:
        text_a_temp = text_a[i1:i2]
        text_b_temp = text_b[j1:j2]
        replacements[text_a_temp] = [i1, i2, text_b_temp]

    return replacements


def get_replacements(text_a, text_b):
    replacements = get_diffs_from_seq(text_a, text_b)
    del_list = []

    for text_a_temp, changes in replacements.items():
        i1, i2, text_b_temp = changes
        if text_a_temp == text_b_temp:
            del_list.append(text_a_temp)

    for key in del_list:
        replacements.pop(key)

    return replacements


def match_without_whitespace(text_a, text_b):
    first_text = text_a
    second_text = text_b
    for item in string.whitespace:
        first_text = first_text.replace(item, "")
        second_text = second_text.replace(item, "")

    return first_text in second_text or second_text in first_text


def find_transposes(text_list_a, text_list_b=None):

    if text_list_b is None:
        text_list_b = text_list_a

    tranposes = []

    for item_a in text_list_a:
        if item_a[0] == "-":
            for item_b in text_list_b:
                if item_b[0] == "+" and match_without_whitespace(item_a[2:], item_b[2:]):
                    tranposes.append(item_a)

    return tranposes


# def get_diffs(text_list_a, text_list_b):
#     diffs = list(ndiff(text_list_a, text_list_b))
#
#     transposed_lines = find_transposes(diffs)
#
#     diff_text = []
#     for diff in diffs:
#
#         if diff in transposed_lines or diff[0] == "?" or diff[0] == " ":
#             continue
#
#         diff_text.append(diff[2:])
#
#     return diff_text
#
#
# def get_diffs_2(text_list_a, text_list_b):
#     diffs = list(ndiff(text_list_a, text_list_b))
#
#     transposed_lines = find_transposes(diffs)
#
#     diff_text = []
#     for diff in diffs:
#
#         if diff[0] == "?" or diff[0] == " ":
#             continue
#
#         diff_text.append(diff)
#
#     return diff_text


def get_diffs(text_list_a, text_list_b):
    diffs = list(ndiff(text_list_a, text_list_b))

    need_replacing = ""

    diff_dict = {}
    for i in range(len(diffs)):
        diff = diffs[i]

        if diff[0] == "?" or diff[0] == " ":
            continue

        if diff[0] == "-":  # if removed from the first
            text_a = diff[2:]
            diff_dict[text_a] = {}
            for j in range(i + 1, len(diffs)):  # look until we find
                diff_temp = diffs[j]
                if diff_temp[0] == "-":  # no replacement
                    diff_dict[text_a][text_a] = [0, len(text_a), ""]
                    i = j
                    break
                elif diff_temp[0] == "+":  # being replaced
                    text_b = diff_temp[2:]
                    diff_dict[text_a] = get_replacements(text_a, text_b)
                    i = j
                    break
        elif diff[0] == "+":  # if added in the second and not replacing
            text_b = diff[2:]
            for j in range(i + 1, len(diffs)):  # look until we find
                diff_temp = diffs[j]
                if diff_temp[0] == "-":  # no replacement
                    text_a = diff_temp[2:]
                    if text_a not in diff_dict:
                        diff_dict[text_a] = {}

                    need_replacing += text_b
                    diff_dict[text_a][text_a] = [0, 0, need_replacing]  # append at the beginning
                    need_replacing = ""
                    i = j
                    break
                elif diff_temp[0] == "+":  # being replaced
                    need_replacing += text_b + "\n\n"
                    text_b = diff_temp[2:]
                    i = j
                    break

    if need_replacing != "":
        diff_dict[""] = {"": "\n" + need_replacing}

    return diff_dict


def get_merge(text_list_a, text_list_b):
    diffs = list(ndiff(text_list_a, text_list_b))

    transposed_lines = find_transposes(diffs)

    merged_text = []
    for diff in diffs:

        if diff in transposed_lines or diff[0] == "?":
            continue

        merged_text.append(diff[2:])

    return merged_text


def load_doc(file):
    if isinstance(file, Document):
        document = file
    else:
        document = RetDoc(file)

    paragraphs = list(document.paragraphs)

    text_list = []

    for paragraph in paragraphs:
        text = paragraph.text
        text_list.append(text)

    return text_list


def merge_text(changes):
    merge_dict = {}
    # assuming ops are ordered based on document position
    for section, change_dict in changes.items():
        merged_text = section
        offset = 0
        for key, change_list in change_dict.items():
            i1, i2, new_text = change_list
            merged_text = merged_text[:i1 + offset] + new_text + merged_text[i2 + offset:]
            offset += len(new_text) - (i2 - i1)

        merge_dict[section] = merged_text

    return merge_dict


# diff_list_b has priority
def merge_diffs(diff_list_a, diff_list_b):
    final_changes = {}

    for section, changes_a in diff_list_a.items():

        if section not in diff_list_b:  # no changes in the final doc in this section
            final_changes[section] = changes_a
        else:  # changes in both sections
            changes_b = diff_list_b[section]
            changes_a_list = [x for t, x in changes_a.items()]
            changes_b_list = [x for t, x in changes_b.items()]

            changes_a_list.sort(key=lambda x: x[0])  # guarantees order
            changes_b_list.sort(key=lambda x: x[0])

            nz_changes_a = len(changes_a_list) > 0
            nz_changes_b = len(changes_b_list) > 0

            final_changes[section] = {}
            while nz_changes_a and nz_changes_b:
                i1, i2, change_text_a = changes_a_list[0]
                j1, j2, change_text_b = changes_b_list[0]

                if i1 <= j2 and j1 <= i2:  # if overlapping, go to the second change
                    ref_text = section[j1:j2]
                    final_changes[section][ref_text] = changes_b_list[0]
                    changes_a_list.pop(0)
                    changes_b_list.pop(0)
                elif i1 <= j1 and i2 <= j1:  # template change is first
                    ref_text = section[i1:i2]
                    final_changes[section][ref_text] = changes_a_list[0]
                    changes_a_list.pop(0)
                elif j1 <= i1 and j2 <= i1:  # final doc change is first
                    ref_text = section[j1:j2]
                    final_changes[section][ref_text] = changes_b_list[0]
                    changes_b_list.pop(0)
                else:  # shouldn't happen
                    print("This shouldn't be happening")

                nz_changes_a = len(changes_a_list) > 0
                nz_changes_b = len(changes_b_list) > 0

    return final_changes


# last doc in docs is highest priority
def merge_docs(pivot_doc, final_doc, docs):

    copyfile(pivot_doc, final_doc)

    text_list = [load_doc(doc) for doc in docs]

    text_list_pivot = load_doc(pivot_doc)  # archived

    diff_with_pivot = [get_diffs(text_list_pivot, x) for x in text_list]

    # collect all the merge changes, giving priority to later items
    final_changes = {}
    for i, diff in enumerate(diff_with_pivot):
        if i == 0:
            final_changes = merge_diffs(diff_with_pivot[0], diff_with_pivot[1])
        else:
            final_changes = merge_diffs(final_changes, diff_with_pivot[i])

    merged_changes = merge_text(final_changes)

    document4 = RetDoc(final_doc)

    for paragraph in list(document4.paragraphs):
        paragraph_text = paragraph.text
        if paragraph_text is '' or paragraph_text not in merged_changes:
            continue
        paragraph.text = merged_changes[paragraph_text]

    document4.save(final_doc)


def replace_doc_text(file, replacements):
    document = Document(file)
    paragraphs = document.paragraphs

    changes = False

    for paragraph in paragraphs:
        text = paragraph.text
        for original, replace in replacements.items():
            if original in replace and replace in text:
                continue
            if original in text:
                changes = True
                text = text.replace(original, replace)
                paragraph.text = text

    if changes:
        print("changing {}".format(file))

    document.save(file)
